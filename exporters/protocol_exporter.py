"""
Экспорт протокола проверки знаний
Заполнение шаблона Protokol_proverki_znanii_OT.docx данными комиссии и работников
"""
import os
import logging
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.export_safe import sanitize_cell_value

logger = logging.getLogger(__name__)


class ProtocolExporter:
    """Экспорт протокола проверки знаний из данных комиссии и записей работников."""

    @staticmethod
    def generate_from_commission(commission_data: dict, protocol_number: str,
                                  worker_records: list, programs_manager,
                                  output_path: str, template_path: str,
                                  data_dir: str) -> tuple[bool, str]:
        """
        Генерация протокола с данными комиссии и работников.

        commission_data — данные комиссии (org_name, order_number, order_date,
                          exam_date, chairman_fio, chairman_position,
                          member1_fio, member2_fio, member3_fio, union_fio)
        protocol_number — номер протокола
        worker_records — список dict с ключами: last_name, first_name, middle_name,
                          snils, position, result, program, date, protocol
        programs_manager — менеджер программ (метод get_program(program_id))
        output_path — путь для сохранения
        template_path — путь к шаблону
        data_dir — директория data

        Возвращает (success: bool, message: str)
        """
        try:
            if not os.path.exists(template_path):
                return False, f"Шаблон протокола не найден:\n{template_path}"

            if not worker_records:
                return False, "Нет записей работников для формирования протокола"

            # Открываем шаблон
            doc = Document(template_path)
            
            # Группируем работников по полному имени
            grouped_workers = ProtocolExporter._group_workers_by_name(worker_records)
            
            ProtocolExporter._fill_worker_table(doc, grouped_workers, programs_manager)
            ProtocolExporter._fill_commission_data(doc, commission_data, protocol_number, programs_manager, worker_records, grouped_workers)

            # Применяем шрифт Times New Roman ко всему документу
            from docx.shared import Pt
            from docx.oxml.ns import qn
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            
            # Для параграфов
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
            
            # Для таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'

            doc.save(output_path)
            return True, f"Протокол сформирован.\nФайл сохранён: {output_path}\nЗаписей: {len(grouped_workers)}"

        except (OSError, ValueError, KeyError) as e:
            logger.error(f"Ошибка генерации протокола: {e}", exc_info=True)
            return False, f"Ошибка формирования протокола: {e}"

    @staticmethod
    def _fill_commission_data(doc, commission_data: dict, protocol_number: str, programs_manager, worker_records: list, grouped_workers: list = None):
        """Заполняет данные комиссии в документе."""
        # Формируем данные программ для плейсхолдеров
        all_programs = ProtocolExporter._format_programs_for_protocol(worker_records, programs_manager)

        # Формируем общий текст всех программ для объединенного плейсхолдера
        all_programs_text = ""
        for i in range(1, 6):
            if i in all_programs:
                prog = all_programs[i]
                if prog.get('hours'):
                    if all_programs_text:
                        all_programs_text += ";"
                    all_programs_text += prog['hours']
        # Режим «одна программа В»: схлопываем программы 6-29 в одну запись
        single_b = {}
        try:
            if programs_manager is not None and hasattr(programs_manager, 'get_single_b_settings'):
                single_b = programs_manager.get_single_b_settings() or {}
        except Exception:
            single_b = {}
        if single_b.get('single_b_mode'):
            sb_hours = (single_b.get('single_b_hours') or '').strip()
            sb_doc = (single_b.get('single_b_doc') or '').strip()
            if sb_hours or sb_doc:
                rebuilt = ProtocolExporter._build_programs_text(worker_records, programs_manager, sb_hours, sb_doc)
                if rebuilt:
                    all_programs_text = rebuilt
        # Добавляем общий текст программ в all_programs для использования
        all_programs['all_programs_text'] = all_programs_text

        # Собираем все регистрационные номера
        all_reg_nums = []
        for i in range(1, 6):
            if i in all_programs:
                base_no = all_programs[i].get('base_no', '')
                if base_no:
                    for bn in base_no.split(';'):
                        if bn and bn not in all_reg_nums:
                            all_reg_nums.append(bn)

        all_reg_nums_text = ';'.join(all_reg_nums)
        all_programs['all_reg_nums'] = all_reg_nums_text

        

# Заполняем текстовые поля по меткам
        for paragraph in doc.paragraphs:
            text = paragraph.text

            # Заменяем плейсхолдеры
            if '{{№ протокола}}' in text:
                text = text.replace('{{№ протокола}}', protocol_number)
            if '{{Название организации}}' in text:
                text = text.replace('{{Название организации}}', commission_data.get('org_name', ''))
            if '{{Дата}}' in text:
                exam_date = commission_data.get('exam_date', '')
                if exam_date:
                    try:
                        from datetime import datetime
                        if isinstance(exam_date, str):
                            if ' ' in exam_date:
                                exam_date = exam_date.split(' ')[0]
                            dt = datetime.strptime(exam_date, '%Y-%m-%d')
                            exam_date = dt.strftime('%d.%m.%Y')
                        elif isinstance(exam_date, datetime):
                            exam_date = exam_date.strftime('%d.%m.%Y')
                    except:
                        pass
                    text = text.replace('{{Дата}}', exam_date)
                else:
                    for rec in worker_records:
                        rec_date = rec.get('date', '')
                        if rec_date:
                            try:
                                from datetime import datetime
                                if ' ' in rec_date:
                                    rec_date = rec_date.split(' ')[0]
                                dt = datetime.strptime(rec_date, '%Y-%m-%d')
                                rec_date = dt.strftime('%d.%m.%Y')
                            except:
                                pass
                            text = text.replace('{{Дата}}', rec_date)
                            break
                    else:
                        text = text.replace('{{Дата}}', '')
            if '{{Приказ о создании комиссии}}' in text:
                order_number = commission_data.get('order_number', '')
                order_date = commission_data.get('order_date', '')
                if order_number and order_date:
                    order_text = f"{order_number} от {order_date}"
                elif order_number:
                    order_text = order_number
                else:
                    order_text = ''
                text = text.replace('{{Приказ о создании комиссии}}', order_text)
            if '{{ФИО председателя}}' in text:
                text = text.replace('{{ФИО председателя}}', commission_data.get('chairman_fio', ''))
            if '{{Должность председателя}}' in text:
                text = text.replace('{{Должность председателя}}', commission_data.get('chairman_position', ''))
            if '{{ФИО члена комиссии №1}}' in text:
                text = text.replace('{{ФИО члена комиссии №1}}', commission_data.get('member1_fio', ''))
            if '{{должность члена комиссии №1}}' in text:
                text = text.replace('{{должность члена комиссии №1}}', commission_data.get('member1_position', ''))
            if '{{ФИО члена комиссии №2}}' in text:
                text = text.replace('{{ФИО члена комиссии №2}}', commission_data.get('member2_fio', ''))
            if '{{должность члена комиссии №2}}' in text:
                text = text.replace('{{должность члена комиссии №2}}', commission_data.get('member2_position', ''))
            if '{{ФИО члена комиссии №3}}' in text:
                text = text.replace('{{ФИО члена комиссии №3}}', commission_data.get('member3_fio', ''))
            if '{{должность члена комиссии №3}}' in text:
                text = text.replace('{{должность члена комиссии №3}}', commission_data.get('member3_position', ''))
            if '{{ФИО представителя профсоюза}}' in text:
                union_fio = commission_data.get('union_fio', '').strip()
                text = text.replace('{{ФИО представителя профсоюза}}', union_fio)
            if '{{ФИО профсоюза}}' in text:
                union_fio = commission_data.get('union_fio', '').strip()
                text = text.replace('{{ФИО профсоюза}}', union_fio)
            if '{{должность представителя профсоюза}}' in text:
                union_pos = commission_data.get('union_position', '').strip()
                text = text.replace('{{должность представителя профсоюза}}', union_pos)
            if '{{Должность профсоюза}}' in text:
                union_pos = commission_data.get('union_position', '').strip()
                text = text.replace('{{Должность профсоюза}}', union_pos)

            # Объединённые плейсхолдеры программ
            if '{{AllPrograms}}' in text:
                text = text.replace('{{AllPrograms}}', all_programs.get('all_programs_text', ''))

            # Результат - всегда "Удовлетворительно"
            if '{{Result}}' in text:
                text = text.replace('{{Result}}', 'Удовлетворительно')

            # Записываем изменённый текст обратно
            paragraph.text = text

            # {{End}} - удаляем
            paragraph.text = paragraph.text.replace('{{End}}', '')

        # Обрабатываем таблицы
        for tbl_idx, tbl in enumerate(doc.tables):
            for row_idx, row in enumerate(tbl.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text
                    reg_nums = all_programs.get('all_reg_nums', '')

                    # AllRegNumbers
                    if '{{AllRegNumbers}}' in cell_text:
                        new_text = cell_text.replace('{{AllRegNumbers}}', reg_nums)
                        cell.text = new_text

                    # Result - подставляем значение по умолчанию
                    if '{{Result}}' in cell_text:
                        result_text = 'Удовлетворительно'
                        new_text = cell_text.replace('{{Result}}', result_text)
                        cell.text = new_text
                    
                    # ФИО профсоюза в таблицах
                    if '{{ФИО представителя профсоюза}}' in cell_text:
                        union_fio = commission_data.get('union_fio', '').strip()
                        new_text = cell_text.replace('{{ФИО представителя профсоюза}}', union_fio)
                        cell.text = new_text
                    if '{{ФИО профсоюза}}' in cell_text:
                        union_fio = commission_data.get('union_fio', '').strip()
                        new_text = cell_text.replace('{{ФИО профсоюза}}', union_fio)
                        cell.text = new_text
                    if '{{должность представителя профсоюза}}' in cell_text:
                        union_pos = commission_data.get('union_position', '').strip()
                        new_text = cell_text.replace('{{должность представителя профсоюза}}', union_pos)
                        cell.text = new_text
                    if '{{Должность профсоюза}}' in cell_text:
                        union_pos = commission_data.get('union_position', '').strip()
                        new_text = cell_text.replace('{{Должность профсоюза}}', union_pos)
                        cell.text = new_text
                    
                    # {{End}} - удаляем
                    if '{{End}}' in cell_text:
                        new_text = cell_text.replace('{{End}}', '')
                        cell.text = new_text

    @staticmethod
    def _group_workers_by_name(worker_records: list) -> list:
        """
        Группирует работников по полному имени (last_name + first_name + middle_name).
        Возвращает список уникальных работников с объединёнными программами.
        """
        from collections import OrderedDict
        grouped = OrderedDict()

        for rec in worker_records:
            full_name = f"{rec.get('last_name', '')} {rec.get('first_name', '')} {rec.get('middle_name', '')}".strip()
            key = full_name.lower()

            if key not in grouped:
                grouped[key] = {
                    'last_name': rec.get('last_name', ''),
                    'first_name': rec.get('first_name', ''),
                    'middle_name': rec.get('middle_name', ''),
                    'full_name': full_name,
                    'snils': rec.get('snils', ''),
                    'position': rec.get('position', ''),
                    'result': rec.get('result', ''),
                    'programs': [],
                    'program_ids': [],
                    'base_nos': []
                }

            # Добавляем программу (название), если ещё не добавлена
            program = rec.get('program', '').strip()
            if program and program not in grouped[key]['programs']:
                grouped[key]['programs'].append(program)

            # Добавляем program_id
            program_id = rec.get('program_id', '').strip()
            if program_id and program_id not in grouped[key]['program_ids']:
                grouped[key]['program_ids'].append(program_id)

            # Добавляем base_no (регистрационный номер)
            base_no = rec.get('base_no', '').strip()
            if base_no and base_no not in grouped[key]['base_nos']:
                grouped[key]['base_nos'].append(base_no)

        # Объединяем program_ids через ";"
        for key in grouped:
            grouped[key]['program_ids_str'] = ';'.join(grouped[key]['program_ids'])
            grouped[key]['base_nos_str'] = ';'.join(grouped[key]['base_nos'])

        return list(grouped.values())

    @staticmethod
    def _format_programs_for_protocol(worker_records: list, programs_manager) -> dict:
        """
        Форматирует программы обучения для заполнения плейсхолдеров в шаблоне протокола.
        Возвращает словарь вида {номер_программы: {'num': '...', 'title': '...', 'hours': '...'}}
        """
        # Собираем все уникальные program_ids из записей
        all_program_ids = set()
        for worker in worker_records:
            # Пробуем разные поля для номера программы
            prog_id = worker.get('program_id', '').strip()
            if prog_id:
                all_program_ids.add(prog_id)
            # Если program_id не найден, пробуем из program (название -> номер)
            if not prog_id:
                program = worker.get('program', '').strip()
                # Пытаемся извлечь номер из названия или использовать как есть
                if program:
                    # Пробуем найти в programs_manager по названию
                    if programs_manager:
                        # Ищем номер программы по названию
                        for pid in range(1, 30):
                            pinfo = programs_manager.get_program(str(pid))
                            if pinfo and pinfo.get('name', '') == program:
                                all_program_ids.add(str(pid))
                                break
                    if not all_program_ids:
                        all_program_ids.add(program)

        # Формируем данные программ для плейсхолдеров
        # sorted_ids - это список номеров программ в порядке (напр. ['2', '12', '22'])
        program_data = {}
        sorted_ids = sorted(all_program_ids, key=lambda x: int(x) if x.isdigit() else 0)

        # Собираем base_nos по каждому program_id
        base_nos_by_prog = {}
        for worker in worker_records:
            prog_id = worker.get('program_id', '').strip()
            base_no = worker.get('base_no', '').strip()
            # Если program_id пустой, используем 'unknown'
            if not prog_id:
                prog_id = 'unknown'
            if base_no:
                if prog_id not in base_nos_by_prog:
                    base_nos_by_prog[prog_id] = []
                if base_no not in base_nos_by_prog[prog_id]:
                    base_nos_by_prog[prog_id].append(base_no)

        for i, prog_id in enumerate(sorted_ids, start=1):
            if i > 5:  # Поддерживаем до 5 программ
                break

            # Get program info if programs_manager is available
            prog_info = {}
            if programs_manager:
                prog_info = programs_manager.get_program(prog_id) or {}

            program_num = prog_id
            # Пробуем получить название из программы (field 'name' или 'title')
            program_title = prog_info.get('name', '') or prog_info.get('title', '') if isinstance(prog_info, dict) else ''
            if not program_title:
                # Если не нашли в programs_manager, используем значение из worker_records
                for worker in worker_records:
                    if worker.get('program_id', '').strip() == prog_id:
                        program_title = worker.get('program', '').strip()
                        if program_title:
                            break

            program_hours = prog_info.get('hours', '') if isinstance(prog_info, dict) else ''
            program_doc = prog_info.get('doc', '') if isinstance(prog_info, dict) else ''

            # Формируем строку в нужном формате
            # Пример: "8 -частовой программе обучения 445 "Оказание первой помощи""
            # Объединяем hours + doc + title в одну строку
            if not program_title and not program_hours and not program_doc:
                # Используем данные из worker_records если programs_manager пустой
                for worker in worker_records:
                    if worker.get('program_id', '').strip() == prog_id:
                        program_title = worker.get('program', '').strip()
                        break

            if program_hours and program_doc and program_title:
                full_program = f"{program_hours} -часовой программе обучения {program_doc} \"{program_title}\""
            elif program_hours and program_doc:
                full_program = f"{program_hours} -часовой программе обучения {program_doc}"
            elif program_hours and program_title:
                full_program = f"{program_hours} -часовой программе обучения \"{program_title}\""
            elif program_hours:
                full_program = f"{program_hours} -часовой программе обучения"
            elif program_doc:
                full_program = f"программе обучения {program_doc}"
            elif program_title:
                # Всегда показываем название программы, даже если hours/doc пустые
                full_program = f"программе обучения \"{program_title}\""
            else:
                full_program = ""

            # Получаем base_no для этого program_id
            base_no_for_prog = base_nos_by_prog.get(prog_id, [])
            base_no_str = ';'.join(base_no_for_prog) if base_no_for_prog else ''

            program_data[i] = {
                'num': program_num,
                'title': program_title,
                'hours': full_program,
                'base_no': base_no_str
            }

        # Также формируем общий список base_nos для всех программ
        all_base_nos = []
        for prog_id in sorted_ids:
            base_nos = base_nos_by_prog.get(prog_id, [])
            for bn in base_nos:
                if bn and bn not in all_base_nos:
                    all_base_nos.append(bn)
        all_base_nos_str = ';'.join(all_base_nos) if all_base_nos else ''

        # Добавляем общий base_no в каждую позицию если пустой
        for i in range(1, len(sorted_ids) + 1):
            if i <= 5 and program_data[i]['base_no'] == '':
                program_data[i]['base_no'] = all_base_nos_str

        return program_data

    @staticmethod
    def _format_single_program_text(prog_id, worker_records: list, programs_manager) -> str:
        """Текст одной программы в формате протокола."""
        prog_info = {}
        if programs_manager:
            prog_info = programs_manager.get_program(prog_id) or {}
        if not isinstance(prog_info, dict):
            prog_info = {}
        title = prog_info.get('name', '') or prog_info.get('title', '')
        if not title:
            for w in worker_records:
                if (w.get('program_id') or '').strip() == prog_id:
                    title = (w.get('program') or '').strip()
                    if title:
                        break
        hours = prog_info.get('hours', '')
        doc = prog_info.get('doc', '')
        if hours and doc and title:
            return f'{hours} -часовой программе обучения {doc} "{title}"'
        if hours and doc:
            return f'{hours} -часовой программе обучения {doc}'
        if hours and title:
            return f'{hours} -часовой программе обучения "{title}"'
        if hours:
            return f'{hours} -часовой программе обучения'
        if doc:
            return f'программе обучения {doc}'
        if title:
            return f'программе обучения "{title}"'
        return ''

    @staticmethod
    def _build_programs_text(worker_records: list, programs_manager, single_b_hours: str = '', single_b_doc: str = '') -> str:
        """Собирает текст программ со схлопыванием программ 6-29 в одну запись."""
        ids = []
        seen = set()
        for w in worker_records:
            pid = (w.get('program_id') or '').strip()
            if pid and pid not in seen:
                seen.add(pid)
                ids.append(pid)

        def _key(x):
            return int(x) if x.isdigit() else 9999
        sorted_ids = sorted(ids, key=_key)

        type_a = [pp for pp in sorted_ids if pp.isdigit() and 1 <= int(pp) <= 5]
        type_b = [pp for pp in sorted_ids if pp.isdigit() and 6 <= int(pp) <= 29]
        other = [pp for pp in sorted_ids if not (pp.isdigit() and 1 <= int(pp) <= 29)]

        entries = []
        for pid in type_a:
            t = ProtocolExporter._format_single_program_text(pid, worker_records, programs_manager)
            if t:
                entries.append(t)

        if type_b:
            hours = (single_b_hours or '').strip()
            doc = (single_b_doc or '').strip()
            tail = 'безопасным методам и приемам выполнения работ повышенной опасности'
            if hours and doc:
                entries.append(f'{hours} -часовой программе обучения № {doc} {tail}')
            elif hours:
                entries.append(f'{hours} -часовой программе обучения {tail}')
            elif doc:
                entries.append(f'программе обучения № {doc} {tail}')

        for pid in other:
            t = ProtocolExporter._format_single_program_text(pid, worker_records, programs_manager)
            if t:
                entries.append(t)

        return ';'.join(e for e in entries if e)

    @staticmethod
    def _fill_worker_table(doc, grouped_workers: list, programs_manager):
        """Заполняет таблицу работников в документе."""
        # Находим таблицу в документе
        table = None
        for tbl_idx, tbl in enumerate(doc.tables):
            if tbl.rows and len(tbl.rows) > 0:
                first_row = tbl.rows[0]
                # Проверяем, есть ли placeholder регистрационных номеров
                has_reg_placeholder = False
                for cell in first_row.cells:
                    if '{{AllRegNumbers}}' in cell.text or 'Регистрационный номер' in cell.text:
                        has_reg_placeholder = True
                        break

                if has_reg_placeholder:
                    # Это таблица с регистрационными номерами - пропускаем
                    continue

                # Проверяем, что это таблица работников по заголовкам
                if len(first_row.cells) >= 8:
                    header_texts = [cell.text.lower() for cell in first_row.cells]
                    if any('ф-и-о' in text or 'фио' in text for text in header_texts):
                        table = tbl
                        break

        if table is None:
            # Если не нашли таблицу по заголовкам, используем вторую таблицу
            if len(doc.tables) > 1:
                table = doc.tables[1]
            elif doc.tables:
                table = doc.tables[0]
            else:
                # Создаем новую таблицу если ее нет
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = '№'
                header_cells[1].text = 'ФИО'
                header_cells[2].text = 'Должность'
                header_cells[3].text = ''
                header_cells[4].text = ''
                header_cells[5].text = 'Результат'
                header_cells[6].text = ''
                header_cells[7].text = ''

        # Очищаем таблицу от существующих данных (оставляем только заголовок)
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        # Заполняем таблицу работников
        for idx, worker in enumerate(grouped_workers):
            row_cells = table.add_row().cells

            # A = порядковый номер
            row_cells[0].text = str(idx + 1)
            # B = ФИО полностью
            row_cells[1].text = sanitize_cell_value(worker['full_name'])
            # C = должность
            row_cells[2].text = sanitize_cell_value(worker.get('position', ''))
            # D = Место работы (оставляем пустым, т.к. это отдельное поле в шаблоне)
            row_cells[3].text = ''
            # E = причина проверки знаний
            row_cells[4].text = 'плановая'
            # F = результат - всегда удовлетворительно
            row_cells[5].text = 'Удовлетворительно'
            # G = регистрационный номер (с новой строки после ;)
            base_nos_str = worker.get('base_nos_str', '')
            if base_nos_str:
                base_nos_str = base_nos_str.replace(';', ';\n')
            row_cells[6].text = sanitize_cell_value(base_nos_str)
            # H = пустое
            row_cells[7].text = ''

            # Форматируем текст в ячейках
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

    @staticmethod
    def export_protocol(records: list, output_path: str, template_path: str, data_dir: str) -> tuple[bool, str]:
        """
        Экспорт протокола из записей журнала (используется в ExamJournalTab).

        records — список JournalRecord с полями: protocol, last_name, first_name, middle_name,
                  snils, position, result, program, date
        output_path — путь для сохранения
        template_path — путь к шаблону
        data_dir — директория data

        Возвращает (success: bool, message: str)
        """
        try:
            if not os.path.exists(template_path):
                return False, f"Шаблон протокола не найден:\n{template_path}"

            if not records:
                return False, "Нет записей для формирования протокола"

            # Загружаем данные комиссии через CommissionManager
            from protocol.commission_manager import CommissionManager
            commission_manager = CommissionManager(data_dir)
            commission_data = commission_manager.load()

            # Получаем ProgramsManager для данных о программах
            programs_manager = None
            try:
                from protocol.programs_manager import ProgramsManager
                programs_manager = ProgramsManager(data_dir)
            except Exception as e:
                logger.warning("Не удалось загрузить ProgramsManager: %s", e, exc_info=True)

            # Преобразуем записи журнала в формат, ожидаемый generate_from_commission
            worker_records = []
            for record in records:
                # Регистрационный номер (base_no)
                base_no = getattr(record, 'base_no', '') or ''
                worker_records.append({
                    'last_name': record.last_name,
                    'first_name': record.first_name,
                    'middle_name': record.middle_name,
                    'snils': record.snils,
                    'position': record.position,
                    'result': record.result,
                    'program': record.program_title,
                    'program_id': record.program_id,
                    'date': record.exam_date,
                    'protocol': record.protocol,
                    'base_no': base_no
                })

            # Используем первый протокол из записей для номера протокола
            protocol_number = records[0].protocol if records else ""

            success, msg = ProtocolExporter.generate_from_commission(
                commission_data=commission_data,
                protocol_number=protocol_number,
                worker_records=worker_records,
                programs_manager=programs_manager,
                output_path=output_path,
                template_path=template_path,
                data_dir=data_dir
            )

            return success, msg

        except (OSError, ValueError, KeyError, AttributeError) as e:
            logger.error(f"Ошибка экспорта протокола из журнала: {e}", exc_info=True)
            return False, f"Ошибка формирования протокола: {e}"